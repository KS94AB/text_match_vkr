from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Iterable, List
import re

from fastapi import HTTPException, UploadFile
from docx import Document

from app.models import DocumentIn

NON_WORD_RE = re.compile(r"[^A-Za-zА-Яа-яЁё0-9]+")


def _safe_doc_id(filename: str, existing: set[str]) -> str:
    stem = Path(filename).stem.lower().replace("ё", "е")
    stem = NON_WORD_RE.sub("_", stem).strip("_") or "document"
    candidate = stem
    counter = 2
    while candidate in existing:
        candidate = f"{stem}_{counter}"
        counter += 1
    existing.add(candidate)
    return candidate


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_txt_text(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="Не удалось декодировать текстовый файл. Используйте UTF-8 или CP1251.")


async def load_documents_from_uploads(files: Iterable[UploadFile]) -> List[DocumentIn]:
    documents: list[DocumentIn] = []
    existing_ids: set[str] = set()

    for upload in files:
        filename = upload.filename or "document"
        suffix = Path(filename).suffix.lower()
        file_bytes = await upload.read()
        if not file_bytes:
            continue

        if suffix == ".docx":
            text = extract_docx_text(file_bytes)
        elif suffix == ".txt":
            text = extract_txt_text(file_bytes)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Файл '{filename}' имеет неподдерживаемый формат. Допустимы .docx и .txt.",
            )

        if not text.strip():
            raise HTTPException(status_code=400, detail=f"Файл '{filename}' не содержит извлекаемого текста.")

        documents.append(
            DocumentIn(
                doc_id=_safe_doc_id(filename, existing_ids),
                title=Path(filename).stem,
                text=text,
            )
        )

    if not documents:
        raise HTTPException(status_code=400, detail="Не удалось получить ни одного документа для анализа.")

    return documents
